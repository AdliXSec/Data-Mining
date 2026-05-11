from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)

IMG = 'images'

def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F3F4')
    shd.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

def add_img(doc, path, width=Inches(5.5), caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.italic = True
        r.font.name = 'Times New Roman'

def set_hdr(cell, color='D6EAF8'):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LAPORAN ANALISIS DATA PREPARATION')
r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Praktikum Data Mining')
r.font.size = Pt(14); r.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Dataset: Teen Mental Health Dataset')
r.font.size = Pt(13); r.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Disusun oleh:\nNaufal Syahruradli')
r.font.size = Pt(12); r.font.name = 'Times New Roman'

for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(datetime.datetime.now().strftime('%B %Y'))
r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# DAFTAR ISI
# ============================================================
doc.add_heading('DAFTAR ISI', level=1)
toc = [
    '1. Pendahuluan',
    '2. Data Understanding',
    '   2.1 Loading Data',
    '   2.2 Karakteristik Data',
    '   2.3 Statistik Deskriptif',
    '   2.4 Visualisasi Data',
    '3. Data Cleaning',
    '   3.1 Pengecekan Missing Values',
    '   3.2 Pengecekan Data Duplikat',
    '   3.3 Pengecekan Data Inkonsisten',
    '   3.4 Deteksi Outlier',
    '4. Data Transformation',
    '   4.1 Data Encoding',
    '   4.2 Standardisasi (Z-Score)',
    '5. Correlation Analysis',
    '6. Analisis Temuan & Kekurangan',
    '7. Kesimpulan',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. PENDAHULUAN
# ============================================================
doc.add_heading('1. Pendahuluan', level=1)
doc.add_paragraph(
    'Laporan ini merupakan dokumentasi analisis tahapan Data Preparation pada dataset '
    '"Teen Mental Health Dataset" sebagai bagian dari tugas Praktikum Data Mining. '
    'Data preparation merupakan tahapan krusial dalam proses data mining karena kualitas '
    'data yang digunakan akan sangat mempengaruhi hasil analisis atau prediksi yang dibuat.')
doc.add_paragraph(
    'Tahapan data preparation yang dibahas meliputi: Data Understanding, Data Cleaning, '
    'dan Data Transformation. Setiap tahapan dianalisis berdasarkan implementasi pada '
    'notebook (code.ipynb), termasuk identifikasi kekurangan dan rekomendasi perbaikan '
    'berdasarkan materi Modul 1 - Data Preparation.')

doc.add_heading('1.1 Deskripsi Dataset', level=2)
doc.add_paragraph(
    'Dataset yang digunakan adalah "Teen_Mental_Health_Dataset.csv" yang berisi data '
    'kesehatan mental remaja. Dataset terdiri dari 1.200 baris dan 13 kolom. '
    'Target variabel: depression_label (0 = Tidak Depresi, 1 = Depresi).')

doc.add_heading('1.2 Deskripsi Variabel', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, t in enumerate(['No', 'Variabel', 'Tipe Data', 'Keterangan']):
    c = table.rows[0].cells[i]; c.text = ''
    r = c.paragraphs[0].add_run(t); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    set_hdr(c)
vars_list = [
    ('1','age','int64','Usia remaja (13-19 tahun)'),
    ('2','gender','object','Jenis kelamin (male/female)'),
    ('3','daily_social_media_hours','float64','Jam penggunaan medsos per hari'),
    ('4','platform_usage','object','Platform (Instagram/TikTok/Both)'),
    ('5','sleep_hours','float64','Jam tidur per hari'),
    ('6','screen_time_before_sleep','float64','Waktu layar sebelum tidur (jam)'),
    ('7','academic_performance','float64','Performa akademik'),
    ('8','physical_activity','float64','Aktivitas fisik (jam)'),
    ('9','social_interaction_level','object','Level interaksi sosial (low/medium/high)'),
    ('10','stress_level','int64','Tingkat stres (0-10)'),
    ('11','anxiety_level','int64','Tingkat kecemasan (0-10)'),
    ('12','addiction_level','int64','Tingkat kecanduan (0-10)'),
    ('13','depression_label','int64','Label depresi (0/1) — TARGET'),
]
for v in vars_list:
    row = table.add_row()
    for i, t in enumerate(v):
        row.cells[i].text = ''
        r = row.cells[i].paragraphs[0].add_run(t); r.font.size = Pt(10); r.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# 2. DATA UNDERSTANDING
# ============================================================
doc.add_heading('2. Data Understanding', level=1)
doc.add_paragraph(
    'Tahap data understanding bertujuan untuk memahami struktur, karakteristik, dan '
    'kualitas data sebelum proses pembersihan dan transformasi.')

doc.add_heading('2.1 Loading Data', level=2)
doc.add_paragraph('Data dimuat menggunakan library Pandas:')
add_code(doc, "import pandas as pd\nimport seaborn as sns\nimport matplotlib.pyplot as plt\n\ndf = pd.read_csv('Teen_Mental_Health_Dataset.csv')")

doc.add_heading('2.2 Karakteristik Data', level=2)
doc.add_paragraph('Berikut adalah tampilan 10 baris pertama dataset menggunakan df.head(10):')
add_img(doc, f'{IMG}/01_head.png', Inches(6), 'Gambar 2.1 — 10 Baris Pertama Dataset')

doc.add_paragraph('Informasi dasar dataset diperoleh menggunakan df.info():')
add_img(doc, f'{IMG}/02_info.png', Inches(4.5), 'Gambar 2.2 — Informasi Dataset (df.info())')
doc.add_paragraph(
    'Dari output di atas, terlihat bahwa seluruh 13 kolom memiliki 1.200 non-null values, '
    'mengindikasikan tidak ada missing values. Tipe data terdiri dari 5 kolom float64, '
    '5 kolom int64, dan 3 kolom object.')

doc.add_heading('2.3 Statistik Deskriptif', level=2)
doc.add_paragraph('Ringkasan statistik deskriptif data numerik menggunakan df.describe():')
add_img(doc, f'{IMG}/03_describe.png', Inches(6), 'Gambar 2.3 — Statistik Deskriptif Data Numerik')

doc.add_paragraph(
    'Dari tabel di atas dapat diketahui bahwa rata-rata usia remaja dalam dataset adalah '
    '15.93 tahun. Rata-rata penggunaan media sosial harian adalah 4.54 jam dengan '
    'rata-rata jam tidur 6.45 jam. Tingkat stress, anxiety, dan addiction memiliki '
    'rata-rata yang relatif serupa (sekitar 5.4-5.5 dari skala 10).')

doc.add_paragraph(
    'Untuk data kategorikal, dari df.describe(include="object") diperoleh:\n'
    '• gender: 2 kategori (male: 615, female: 585)\n'
    '• platform_usage: 3 kategori (Instagram: 411 sebagai modus)\n'
    '• social_interaction_level: 3 kategori (medium: 416 sebagai modus)')

doc.add_page_break()

doc.add_heading('2.4 Visualisasi Data', level=2)
doc.add_heading('a. Distribusi Variabel Numerik (Histogram)', level=3)
doc.add_paragraph(
    'Histogram digunakan untuk memahami distribusi dan penyebaran data numerik:')
add_img(doc, f'{IMG}/04_dist_numerical.png', Inches(6), 'Gambar 2.4 — Distribusi Variabel Numerik')
doc.add_paragraph(
    'Dari histogram di atas terlihat bahwa variabel age memiliki distribusi relatif merata. '
    'Variabel stress_level, anxiety_level, dan addiction_level memiliki distribusi yang '
    'cenderung uniform. Depression_label menunjukkan distribusi biner (0 dan 1).')

doc.add_heading('b. Distribusi Variabel Kategorikal (Bar Plot)', level=3)
doc.add_paragraph('Bar plot menampilkan frekuensi setiap kategori pada variabel kategorikal:')
add_img(doc, f'{IMG}/05_barplot_categorical.png', Inches(5.5), 'Gambar 2.5 — Bar Plot Variabel Kategorikal')
doc.add_paragraph(
    'Dari bar plot terlihat bahwa distribusi gender relatif seimbang (male: 615, female: 585). '
    'Platform Instagram paling banyak digunakan (411), diikuti Both (400) dan TikTok (389). '
    'Level interaksi sosial medium mendominasi (416).')

doc.add_heading('c. Distribusi Variabel Kategorikal (Pie Chart)', level=3)
doc.add_paragraph('Pie chart menunjukkan proporsi relatif setiap kategori:')
add_img(doc, f'{IMG}/06_piechart_categorical.png', Inches(5.5), 'Gambar 2.6 — Pie Chart Variabel Kategorikal')

doc.add_page_break()

# ============================================================
# 3. DATA CLEANING
# ============================================================
doc.add_heading('3. Data Cleaning', level=1)
doc.add_paragraph(
    'Data cleaning adalah proses identifikasi dan perbaikan data yang tidak akurat, rusak, '
    'tidak lengkap, duplikat, atau tidak relevan dari dataset.')

doc.add_heading('3.1 Pengecekan Missing Values', level=2)
doc.add_paragraph('Pengecekan missing values dilakukan secara numerik dan visual:')
add_code(doc, 'print("Jumlah Missing Values:\\n", df.isnull().sum())')
doc.add_paragraph(
    'Hasil: Seluruh kolom memiliki 0 missing values. Visualisasi heatmap berikut '
    'mengkonfirmasi tidak adanya missing values (seluruh area berwarna seragam):')
add_img(doc, f'{IMG}/08_missing_heatmap.png', Inches(5.5), 'Gambar 3.1 — Heatmap Missing Values')
doc.add_paragraph(
    'Dari heatmap di atas, seluruh area berwarna seragam yang mengindikasikan '
    'tidak ada missing values pada dataset. Dataset sudah lengkap.')

doc.add_heading('3.2 Pengecekan Data Duplikat', level=2)
add_code(doc, 'duplicate_count = df.duplicated().sum()\nprint(f"Jumlah Data Duplikat: {duplicate_count}")')
doc.add_paragraph(
    'Hasil: Jumlah data duplikat = 0. Tidak ditemukan baris identik dalam dataset.')

doc.add_heading('3.3 Pengecekan Data Inkonsisten', level=2)
doc.add_paragraph(
    'Verifikasi konsistensi data kategorikal dilakukan untuk memastikan tidak ada '
    'variasi format penulisan. Hasil pengecekan:\n'
    '• gender: hanya "male" dan "female" (konsisten lowercase)\n'
    '• platform_usage: "Instagram", "TikTok", "Both" (konsisten)\n'
    '• social_interaction_level: "low", "medium", "high" (konsisten lowercase)')

doc.add_heading('3.4 Deteksi Outlier', level=2)
doc.add_paragraph('Deteksi outlier dilakukan menggunakan metode IQR dan divisualisasikan dengan boxplot:')
add_code(doc,
    'for col in numerical_cols:\n'
    '    Q1 = df[col].quantile(0.25)\n'
    '    Q3 = df[col].quantile(0.75)\n'
    '    IQR = Q3 - Q1\n'
    '    lower = Q1 - 1.5 * IQR\n'
    '    upper = Q3 + 1.5 * IQR\n'
    '    outliers = df[(df[col] < lower) | (df[col] > upper)]')

add_img(doc, f'{IMG}/07_boxplot.png', Inches(6), 'Gambar 3.2 — Boxplot untuk Deteksi Outlier')
doc.add_paragraph(
    'Dari boxplot di atas, tidak terlihat adanya titik data di luar whisker pada seluruh '
    'variabel numerik. Hal ini mengkonfirmasi hasil deteksi IQR bahwa tidak ditemukan '
    'outlier pada dataset ini. Semua nilai berada dalam batas wajar.')

doc.add_page_break()

# ============================================================
# 4. DATA TRANSFORMATION
# ============================================================
doc.add_heading('4. Data Transformation', level=1)
doc.add_paragraph(
    'Transformasi data bertujuan untuk mengubah data mentah menjadi format yang sesuai '
    'untuk analisis dan pemodelan machine learning.')

doc.add_heading('4.1 Data Encoding', level=2)
doc.add_paragraph(
    'Encoding dilakukan untuk mengubah variabel kategorikal menjadi format numerik:')

doc.add_heading('a. Label Encoding — gender', level=3)
add_code(doc,
    'from sklearn.preprocessing import LabelEncoder\n'
    'le_gender = LabelEncoder()\n'
    "df['gender'] = le_gender.fit_transform(df['gender'])")
doc.add_paragraph(
    'Hasil: female → 0, male → 1. Label Encoding dipilih karena gender bersifat biner.')

doc.add_heading('b. Ordinal Mapping — social_interaction_level', level=3)
add_code(doc,
    "social_mapping = {'low': 0, 'medium': 1, 'high': 2}\n"
    "df['social_interaction_level'] = df['social_interaction_level'].map(social_mapping)")
doc.add_paragraph(
    'Mapping manual karena variabel bersifat ordinal (berurutan): low → 0, medium → 1, high → 2.')

doc.add_heading('c. One-Hot Encoding — platform_usage', level=3)
add_code(doc, "df = pd.get_dummies(df, columns=['platform_usage'], prefix='platform')")
doc.add_paragraph(
    'One-Hot Encoding dipilih karena platform_usage bersifat nominal (tidak ada urutan). '
    'Menghasilkan 3 kolom biner: platform_Both, platform_Instagram, platform_TikTok.')

doc.add_heading('4.2 Standardisasi (Z-Score)', level=2)
doc.add_paragraph('Standardisasi Z-Score diterapkan pada variabel numerik. Rumus: Z = (X - μ) / σ')
add_code(doc,
    "from sklearn.preprocessing import StandardScaler\n"
    "scaler = StandardScaler()\n"
    "df[features_to_scale] = scaler.fit_transform(df[features_to_scale])")
doc.add_paragraph(
    'Setelah standardisasi, setiap fitur memiliki mean ≈ 0 dan standar deviasi ≈ 1. '
    'Penting untuk algoritma yang sensitif terhadap skala seperti Logistic Regression dan SVM.')

doc.add_paragraph('Berikut tampilan data setelah seluruh proses encoding dan standardisasi:')
add_img(doc, f'{IMG}/09_after_transform.png', Inches(6), 'Gambar 4.1 — Data Setelah Encoding & Standardisasi')

doc.add_page_break()

# ============================================================
# 5. CORRELATION ANALYSIS
# ============================================================
doc.add_heading('5. Correlation Analysis', level=1)
doc.add_paragraph('Analisis korelasi dilakukan untuk melihat hubungan antar variabel:')
add_code(doc,
    "plt.figure(figsize=(13, 10))\n"
    "sns.heatmap(df.corr(), annot=True, cmap='coolwarm', fmt='.2f')\n"
    "plt.title('Correlation Matrix after Preprocessing')\n"
    "plt.show()")
add_img(doc, f'{IMG}/10_correlation_heatmap.png', Inches(5.5), 'Gambar 5.1 — Heatmap Korelasi')
doc.add_paragraph(
    'Dari heatmap korelasi di atas, dapat diamati hubungan antar variabel. '
    'Korelasi positif ditandai warna merah dan negatif berwarna biru. '
    'Analisis ini membantu mengidentifikasi fitur yang berkorelasi dengan target '
    '(depression_label) serta mendeteksi kemungkinan multikolinearitas antar fitur.')

doc.add_page_break()

# ============================================================
# 6. ANALISIS TEMUAN & KEKURANGAN
# ============================================================
doc.add_heading('6. Analisis Temuan dan Kekurangan', level=1)
doc.add_paragraph(
    'Berdasarkan review implementasi data preparation pada notebook code.ipynb '
    'dan materi Modul 1, ditemukan beberapa temuan:')

doc.add_heading('6.1 Urutan Proses yang Perlu Diperbaiki', level=2)
doc.add_paragraph(
    'Pada notebook, urutan proses adalah: Data Understanding → Data Preprocessing '
    '(Encoding + Scaling) → Correlation → Data Cleaning. Idealnya data cleaning '
    'dilakukan sebelum transformasi agar proses transformasi menggunakan data yang sudah bersih.')
doc.add_paragraph('Urutan yang seharusnya:')
doc.add_paragraph('1. Data Understanding', style='List Number')
doc.add_paragraph('2. Data Cleaning (missing values, duplikat, outlier)', style='List Number')
doc.add_paragraph('3. Data Transformation (encoding, scaling)', style='List Number')
doc.add_paragraph('4. Correlation Analysis', style='List Number')

doc.add_heading('6.2 Visualisasi yang Perlu Dilengkapi', level=2)
doc.add_paragraph(
    'Notebook awal hanya menggunakan displot untuk semua kolom. Berdasarkan Modul 1:\n'
    '• Variabel kuantitatif: Histogram (sudah ada) + Boxplot (perlu ditambah)\n'
    '• Variabel kualitatif: Bar Plot / Pie Chart (perlu ditambah)\n\n'
    'Pada laporan ini, visualisasi sudah dilengkapi dengan boxplot, bar plot, dan pie chart.')

doc.add_heading('6.3 Validasi Invalid Values', level=2)
doc.add_paragraph(
    'Modul 1 membahas koreksi nilai tidak valid. Meskipun dataset ini tidak memiliki '
    'invalid values, proses verifikasi perlu ditunjukkan (misalnya: memastikan usia '
    'dalam range 13-19 tahun, jam tidur 0-24 jam, dll).')

doc.add_heading('6.4 Visualisasi Missing Values', level=2)
doc.add_paragraph(
    'Notebook awal hanya mengecek isnull().sum() secara numerik. Modul 1 menyarankan '
    'penggunaan heatmap untuk visualisasi missing values. Pada laporan ini sudah '
    'ditambahkan heatmap missing values (Gambar 3.1).')

doc.add_heading('6.5 Penggunaan head() dan tail()', level=2)
doc.add_paragraph(
    'Modul 1 menyebutkan head() dan tail() sebagai langkah dasar data understanding. '
    'Notebook menampilkan df secara keseluruhan tanpa head()/tail() eksplisit.')

doc.add_page_break()

# ============================================================
# 7. KESIMPULAN
# ============================================================
doc.add_heading('7. Kesimpulan', level=1)
doc.add_paragraph(
    'Berdasarkan analisis yang telah dilakukan, tahapan data preparation pada dataset '
    'Teen Mental Health sudah mencakup komponen-komponen berikut:')
doc.add_paragraph('• Data Understanding: loading data, eksplorasi karakteristik, statistik deskriptif, visualisasi')
doc.add_paragraph('• Data Cleaning: pengecekan missing values (0), duplikat (0), dan deteksi outlier (0)')
doc.add_paragraph('• Data Transformation: Label Encoding, Ordinal Mapping, One-Hot Encoding, dan Standardisasi Z-Score')
doc.add_paragraph('• Correlation Analysis: heatmap korelasi antar variabel')

doc.add_paragraph()
doc.add_paragraph('Rekomendasi perbaikan yang disarankan:')
doc.add_paragraph('1. Perbaiki urutan proses: Data Cleaning dilakukan sebelum Data Transformation.', style='List Number')
doc.add_paragraph('2. Lengkapi visualisasi dengan boxplot untuk deteksi outlier visual dan bar/pie chart untuk variabel kategorik.', style='List Number')
doc.add_paragraph('3. Tambahkan validasi nilai tidak valid sebagai langkah verifikasi.', style='List Number')
doc.add_paragraph('4. Tambahkan visualisasi heatmap missing values.', style='List Number')
doc.add_paragraph('5. Gunakan head() dan tail() untuk preview data.', style='List Number')
doc.add_paragraph('6. Perkaya narasi markdown untuk dokumentasi yang lebih baik.', style='List Number')

doc.add_paragraph()
doc.add_paragraph(
    'Secara keseluruhan, dataset Teen Mental Health tergolong "bersih" — tidak ada '
    'missing values, data duplikat, maupun outlier — sehingga proses cleaning bersifat '
    'verifikasi. Proses transformasi (encoding dan standardisasi) sudah tepat dan data '
    'siap digunakan untuk tahap selanjutnya yaitu Data Modeling.')

# ============================================================
# SAVE
# ============================================================
doc.save('laporan_v2.docx')
print("laporan_v2.docx berhasil dibuat dengan gambar!")
