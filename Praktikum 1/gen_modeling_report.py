import json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

with open('modeling_results.json') as f:
    R = json.load(f)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54); s.left_margin=Cm(3); s.right_margin=Cm(2.54)
st = doc.styles['Normal']; st.font.name='Times New Roman'; st.font.size=Pt(12)
st.paragraph_format.line_spacing=1.5
for lv in range(1,4):
    h=doc.styles[f'Heading {lv}']; h.font.name='Times New Roman'; h.font.color.rgb=RGBColor(0,0,0)
    h.font.size=Pt(16 if lv==1 else 14 if lv==2 else 12)

IMG='images_modeling'

def code(doc,t):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(9.5)
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'F2F3F4'); sh.set(qn('w:val'),'clear')
    p.paragraph_format.element.get_or_add_pPr().append(sh)

def img(doc,path,w=Inches(5.5),cap=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path,width=w)
    if cap:
        c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=c.add_run(cap); r.font.size=Pt(10); r.italic=True; r.font.name='Times New Roman'

def shd(cell,color='D6EAF8'):
    s=OxmlElement('w:shd'); s.set(qn('w:fill'),color); s.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(s)

def mk_table(doc, headers, rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10); r.font.name='Times New Roman'
        shd(c)
    for row in rows:
        rw=t.add_row()
        for i,v in enumerate(row):
            rw.cells[i].text=''
            r=rw.cells[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(10); r.font.name='Times New Roman'

# ===== COVER =====
for _ in range(6): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('LAPORAN ANALISIS DATA MODELING'); r.bold=True; r.font.size=Pt(18); r.font.name='Times New Roman'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Praktikum Data Mining\nDataset: Teen Mental Health Dataset'); r.font.size=Pt(14)
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Disusun oleh:\nNaufal Syahruradli'); r.font.size=Pt(12)
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run(datetime.datetime.now().strftime('%B %Y')); r.font.size=Pt(12)
doc.add_page_break()

# ===== DAFTAR ISI =====
doc.add_heading('DAFTAR ISI',level=1)
for item in ['1. Pendahuluan','2. Train/Test Split','3. Model 1: Random Forest Classifier',
    '4. Model 2: Logistic Regression','5. Model 3: Support Vector Machine (SVM)',
    '6. Perbandingan Model','7. Feature Importance','8. Analisis Overfitting/Underfitting',
    '9. Kesimpulan']:
    doc.add_paragraph(item).paragraph_format.space_after=Pt(2)
doc.add_page_break()

# ===== 1. PENDAHULUAN =====
doc.add_heading('1. Pendahuluan',level=1)
doc.add_paragraph('Laporan ini mendokumentasikan tahapan Data Modeling pada dataset "Teen Mental Health Dataset". Data modeling adalah tahapan utama dalam proses data mining di mana teknik machine learning diterapkan pada data yang telah dipreparasi untuk membangun model prediktif.')
doc.add_paragraph('Pada tahap ini, tiga algoritma klasifikasi digunakan untuk memprediksi label depresi (depression_label) pada remaja, yaitu: Random Forest Classifier, Logistic Regression, dan Support Vector Machine (SVM). Setiap model dievaluasi menggunakan metrik Accuracy, Precision, Recall, dan F1-Score.')

doc.add_heading('1.1 Rumus Metrik Evaluasi',level=2)
doc.add_paragraph('Berikut adalah rumus metrik evaluasi yang digunakan dalam laporan ini:')
doc.add_paragraph('a. Accuracy — mengukur proporsi prediksi benar dari keseluruhan data:')
code(doc,'Accuracy = (TP + TN) / (TP + TN + FP + FN)')
doc.add_paragraph('b. Precision — mengukur ketepatan prediksi positif:')
code(doc,'Precision = TP / (TP + FP)')
doc.add_paragraph('c. Recall (Sensitivity) — mengukur kemampuan model menemukan semua data positif:')
code(doc,'Recall = TP / (TP + FN)')
doc.add_paragraph('d. F1-Score — rata-rata harmonik antara Precision dan Recall:')
code(doc,'F1-Score = 2 × (Precision × Recall) / (Precision + Recall)')
doc.add_paragraph('Keterangan:\n• TP (True Positive): data positif yang diprediksi benar sebagai positif\n• TN (True Negative): data negatif yang diprediksi benar sebagai negatif\n• FP (False Positive): data negatif yang salah diprediksi sebagai positif\n• FN (False Negative): data positif yang salah diprediksi sebagai negatif')
doc.add_page_break()

# ===== 2. TRAIN/TEST SPLIT =====
doc.add_heading('2. Train/Test Split',level=1)
doc.add_paragraph(f'Dataset dibagi menjadi data training (80%) dan data testing (20%) menggunakan stratified sampling untuk menjaga distribusi kelas target yang seimbang.')
doc.add_heading('2.1 Kode Implementasi',level=2)
code(doc,"from sklearn.model_selection import train_test_split\n\n# Memisahkan fitur (X) dan target (y)\nX = df.drop('depression_label', axis=1)\ny = df['depression_label']\n\n# Split data: 80% training, 20% testing\nX_train, X_test, y_train, y_test = train_test_split(\n    X, y, test_size=0.2, random_state=42, stratify=y)")

doc.add_heading('2.2 Penjelasan Kode',level=2)
doc.add_paragraph('• df.drop(\'depression_label\', axis=1): Memisahkan fitur (X) dengan menghapus kolom target dari dataframe.')
doc.add_paragraph('• test_size=0.2: Mengalokasikan 20% data untuk testing dan 80% untuk training.')
doc.add_paragraph('• random_state=42: Menetapkan seed agar hasil split konsisten dan reproducible.')
doc.add_paragraph('• stratify=y: Memastikan proporsi kelas target (0 dan 1) tetap sama di data training dan testing.')

doc.add_heading('2.3 Hasil Split Data',level=2)
sp=R['split']
mk_table(doc,['Keterangan','Jumlah','Kelas 0','Kelas 1'],[
    ['Data Training (80%)',sp['train'],sp['train_0'],sp['train_1']],
    ['Data Testing (20%)',sp['test'],sp['test_0'],sp['test_1']],
    ['Total',sp['total'],sp['train_0']+sp['test_0'],sp['train_1']+sp['test_1']]
])
doc.add_paragraph(f'\nJumlah fitur yang digunakan: {sp["n_features"]} fitur.')
doc.add_paragraph(f'Fitur: {", ".join(sp["features"])}')
img(doc,f'{IMG}/01_split_distribution.png',Inches(5.5),'Gambar 2.1 — Distribusi Target pada Split Data')
doc.add_paragraph('Dari visualisasi di atas terlihat bahwa dataset memiliki distribusi kelas yang sangat tidak seimbang (imbalanced). Kelas 0 (Tidak Depresi) mendominasi dengan jumlah jauh lebih besar dibandingkan kelas 1 (Depresi). Stratified sampling memastikan proporsi ini terjaga di kedua subset.')
doc.add_page_break()

# ===== 3. RANDOM FOREST =====
doc.add_heading('3. Model 1: Random Forest Classifier',level=1)
doc.add_heading('3.1 Deskripsi Algoritma',level=2)
doc.add_paragraph('Random Forest adalah algoritma ensemble learning yang membangun banyak decision tree selama proses training dan menghasilkan prediksi berdasarkan voting mayoritas dari seluruh tree. Keunggulannya: robust terhadap overfitting, mampu menangani data berdimensi tinggi, dan dapat mengukur feature importance.')

doc.add_heading('3.2 Kode Implementasi',level=2)
code(doc,"from sklearn.ensemble import RandomForestClassifier\nfrom sklearn.metrics import accuracy_score, classification_report, confusion_matrix\n\nrf_model = RandomForestClassifier(\n    n_estimators=100,    # Jumlah pohon keputusan\n    random_state=42,     # Seed untuk reproducibility\n    max_depth=10,        # Kedalaman maksimum pohon\n    min_samples_split=5, # Min sampel untuk split node\n    min_samples_leaf=2   # Min sampel pada leaf node\n)\nrf_model.fit(X_train, y_train)\ny_pred_rf = rf_model.predict(X_test)")

doc.add_heading('3.3 Penjelasan Hyperparameter',level=2)
mk_table(doc,['Parameter','Nilai','Penjelasan'],[
    ['n_estimators','100','Jumlah pohon keputusan yang dibangun dalam forest'],
    ['max_depth','10','Batas kedalaman maksimum setiap pohon untuk mencegah overfitting'],
    ['min_samples_split','5','Jumlah minimum sampel yang diperlukan untuk membagi node internal'],
    ['min_samples_leaf','2','Jumlah minimum sampel yang diperlukan pada leaf node'],
    ['random_state','42','Seed untuk memastikan hasil yang reproducible']
])

doc.add_heading('3.4 Hasil Evaluasi',level=2)
rf=R['Random Forest']
doc.add_paragraph(f'• Training Accuracy: {rf["train_acc"]}\n• Testing Accuracy: {rf["test_acc"]}\n• Precision: {rf["precision"]}\n• Recall: {rf["recall"]}\n• F1-Score: {rf["f1"]}')
doc.add_paragraph(f'\nConfusion Matrix:\n• True Negative (TN): {rf["cm"][0][0]} — data tidak depresi yang diprediksi benar\n• False Positive (FP): {rf["cm"][0][1]} — data tidak depresi yang salah diprediksi depresi\n• False Negative (FN): {rf["cm"][1][0]} — data depresi yang salah diprediksi tidak depresi\n• True Positive (TP): {rf["cm"][1][1]} — data depresi yang diprediksi benar')
doc.add_paragraph('\nAnalisis: Model Random Forest memiliki akurasi tinggi (97.5%) namun gagal mendeteksi kelas positif (Depresi) — Precision, Recall, dan F1-Score bernilai 0. Hal ini disebabkan oleh ketidakseimbangan kelas yang sangat ekstrem dalam dataset.')
doc.add_page_break()

# ===== 4. LOGISTIC REGRESSION =====
doc.add_heading('4. Model 2: Logistic Regression',level=1)
doc.add_heading('4.1 Deskripsi Algoritma',level=2)
doc.add_paragraph('Logistic Regression adalah algoritma klasifikasi yang memodelkan probabilitas kelas menggunakan fungsi logistik (sigmoid). Meskipun namanya mengandung "regression", algoritma ini digunakan untuk klasifikasi biner atau multi-kelas. Rumus fungsi sigmoid: σ(z) = 1 / (1 + e^(-z)), di mana z = w₁x₁ + w₂x₂ + ... + wₙxₙ + b.')

doc.add_heading('4.2 Kode Implementasi',level=2)
code(doc,"from sklearn.linear_model import LogisticRegression\n\nlr_model = LogisticRegression(\n    random_state=42,  # Seed untuk reproducibility\n    max_iter=1000,    # Iterasi maksimum untuk konvergensi\n    solver='lbfgs'    # Algoritma optimasi\n)\nlr_model.fit(X_train, y_train)\ny_pred_lr = lr_model.predict(X_test)")

doc.add_heading('4.3 Penjelasan Hyperparameter',level=2)
mk_table(doc,['Parameter','Nilai','Penjelasan'],[
    ['max_iter','1000','Jumlah iterasi maksimum agar algoritma mencapai konvergensi'],
    ['solver','lbfgs','Metode optimasi berbasis kuadrat, cocok untuk data besar dengan regularisasi L2'],
    ['random_state','42','Seed untuk memastikan hasil yang reproducible']
])

doc.add_heading('4.4 Hasil Evaluasi',level=2)
lr=R['Logistic Regression']
doc.add_paragraph(f'• Training Accuracy: {lr["train_acc"]}\n• Testing Accuracy: {lr["test_acc"]}\n• Precision: {lr["precision"]}\n• Recall: {lr["recall"]}\n• F1-Score: {lr["f1"]}')
doc.add_paragraph(f'\nConfusion Matrix:\n• True Negative (TN): {lr["cm"][0][0]}\n• False Positive (FP): {lr["cm"][0][1]}\n• False Negative (FN): {lr["cm"][1][0]}\n• True Positive (TP): {lr["cm"][1][1]}')
doc.add_paragraph('\nAnalisis: Logistic Regression menunjukkan performa terbaik di antara ketiga model. Akurasi testing 98.75%, Precision 1.0 (tidak ada false positive), Recall 0.5 (berhasil mendeteksi 3 dari 6 kasus depresi), dan F1-Score 0.6667. Model ini berhasil mengidentifikasi sebagian kasus depresi tanpa menghasilkan false alarm.')
doc.add_page_break()

# ===== 5. SVM =====
doc.add_heading('5. Model 3: Support Vector Machine (SVM)',level=1)
doc.add_heading('5.1 Deskripsi Algoritma',level=2)
doc.add_paragraph('Support Vector Machine (SVM) adalah algoritma klasifikasi yang bekerja dengan mencari hyperplane optimal yang memaksimalkan margin antara dua kelas. SVM menggunakan kernel trick untuk menangani data yang tidak dapat dipisahkan secara linear. Kernel RBF (Radial Basis Function) memetakan data ke dimensi lebih tinggi.')

doc.add_heading('5.2 Kode Implementasi',level=2)
code(doc,"from sklearn.svm import SVC\n\nsvm_model = SVC(\n    kernel='rbf',     # Kernel Radial Basis Function\n    random_state=42,  # Seed untuk reproducibility\n    C=1.0,            # Parameter regularisasi\n    gamma='scale'     # Koefisien kernel\n)\nsvm_model.fit(X_train, y_train)\ny_pred_svm = svm_model.predict(X_test)")

doc.add_heading('5.3 Penjelasan Hyperparameter',level=2)
mk_table(doc,['Parameter','Nilai','Penjelasan'],[
    ['kernel','rbf','Radial Basis Function — memetakan data ke dimensi lebih tinggi'],
    ['C','1.0','Parameter regularisasi; nilai lebih kecil = regularisasi lebih kuat'],
    ['gamma','scale','Koefisien kernel; "scale" menggunakan 1/(n_features × var(X))'],
    ['random_state','42','Seed untuk memastikan hasil yang reproducible']
])

doc.add_heading('5.4 Hasil Evaluasi',level=2)
sv=R['SVM']
doc.add_paragraph(f'• Training Accuracy: {sv["train_acc"]}\n• Testing Accuracy: {sv["test_acc"]}\n• Precision: {sv["precision"]}\n• Recall: {sv["recall"]}\n• F1-Score: {sv["f1"]}')
doc.add_paragraph(f'\nConfusion Matrix:\n• True Negative (TN): {sv["cm"][0][0]}\n• False Positive (FP): {sv["cm"][0][1]}\n• False Negative (FN): {sv["cm"][1][0]}\n• True Positive (TP): {sv["cm"][1][1]}')
doc.add_paragraph('\nAnalisis: SVM memiliki pola yang sama dengan Random Forest — akurasi tinggi (97.5%) namun gagal mendeteksi kelas positif. Seluruh data diprediksi sebagai kelas 0 (Tidak Depresi), menghasilkan Precision, Recall, dan F1-Score bernilai 0 untuk kelas positif.')
doc.add_page_break()

# ===== 6. PERBANDINGAN MODEL =====
doc.add_heading('6. Perbandingan Model',level=1)
doc.add_heading('6.1 Kode Perbandingan',level=2)
code(doc,"from sklearn.metrics import precision_score, recall_score, f1_score\n\nmodel_names = ['Random Forest', 'Logistic Regression', 'SVM']\npredictions = [y_pred_rf, y_pred_lr, y_pred_svm]\n\nfor name, y_pred in zip(model_names, predictions):\n    print(f'{name}: Acc={accuracy_score(y_test, y_pred):.4f}, '\n          f'Prec={precision_score(y_test, y_pred):.4f}, '\n          f'Rec={recall_score(y_test, y_pred):.4f}, '\n          f'F1={f1_score(y_test, y_pred):.4f}')")

doc.add_heading('6.2 Tabel Perbandingan',level=2)
img(doc,f'{IMG}/08_comparison_table.png',Inches(6),'Gambar 6.1 — Tabel Perbandingan Performa Model')

doc.add_heading('6.3 Confusion Matrix Ketiga Model',level=2)
img(doc,f'{IMG}/02_confusion_matrices.png',Inches(6),'Gambar 6.2 — Confusion Matrix Ketiga Model')
doc.add_paragraph('Dari confusion matrix terlihat bahwa:\n• Random Forest dan SVM memprediksi seluruh data sebagai kelas 0 (Tidak Depresi)\n• Logistic Regression berhasil mendeteksi 3 dari 6 kasus depresi (TP=3) tanpa false positive (FP=0)')

doc.add_heading('6.4 Classification Report',level=2)
img(doc,f'{IMG}/03_classification_reports.png',Inches(6),'Gambar 6.3 — Classification Report Ketiga Model')

doc.add_heading('6.5 Visualisasi Perbandingan Metrik',level=2)
img(doc,f'{IMG}/04_model_comparison.png',Inches(5.5),'Gambar 6.4 — Perbandingan Metrik Evaluasi')
doc.add_paragraph('Dari grafik perbandingan, Logistic Regression unggul pada metrik Precision, Recall, dan F1-Score meskipun ketiga model memiliki accuracy yang serupa.')

doc.add_heading('6.6 Perbandingan Akurasi',level=2)
img(doc,f'{IMG}/05_accuracy_comparison.png',Inches(5),'Gambar 6.5 — Perbandingan Akurasi Model')
doc.add_paragraph('Akurasi ketiga model sangat tinggi (>97%), namun angka ini menyesatkan karena dataset sangat imbalanced. Model yang hanya memprediksi semua data sebagai kelas mayoritas pun sudah mendapat akurasi tinggi.')
doc.add_page_break()

# ===== 7. FEATURE IMPORTANCE =====
doc.add_heading('7. Feature Importance',level=1)
doc.add_heading('7.1 Kode Implementasi',level=2)
code(doc,"feature_importance = pd.DataFrame({\n    'Feature': X.columns,\n    'Importance': rf_model.feature_importances_\n}).sort_values('Importance', ascending=False)\n\nprint(feature_importance.to_string(index=False))")
doc.add_paragraph('Feature importance dari Random Forest menunjukkan kontribusi relatif setiap fitur dalam proses pengambilan keputusan model.')

doc.add_heading('7.2 Hasil Feature Importance',level=2)
fi_rows = [[f,str(v)] for f,v in R['feature_importance'].items()]
mk_table(doc,['Fitur','Importance'],fi_rows)
doc.add_paragraph()
img(doc,f'{IMG}/06_feature_importance.png',Inches(5.5),'Gambar 7.1 — Feature Importance Random Forest')
doc.add_paragraph('Interpretasi:\n• sleep_hours (0.2099) adalah fitur paling berpengaruh — jam tidur merupakan faktor utama dalam prediksi depresi remaja\n• daily_social_media_hours (0.1919) menjadi faktor kedua terpenting — penggunaan media sosial berkorelasi kuat dengan depresi\n• stress_level (0.1360) dan anxiety_level (0.1200) juga memiliki pengaruh signifikan\n• Platform yang digunakan (Both/Instagram/TikTok) memiliki pengaruh paling kecil')
doc.add_page_break()

# ===== 8. OVERFITTING =====
doc.add_heading('8. Analisis Overfitting/Underfitting',level=1)
doc.add_paragraph('Overfitting terjadi ketika model terlalu mempelajari data training sehingga performanya buruk pada data baru. Underfitting terjadi ketika model terlalu sederhana untuk menangkap pola data. Untuk mengecek, dibandingkan akurasi training vs testing:')
code(doc,"# Cek overfitting\nfor name, model in models.items():\n    train_acc = accuracy_score(y_train, model.predict(X_train))\n    test_acc = accuracy_score(y_test, model.predict(X_test))\n    print(f'{name}: Train={train_acc:.4f}, Test={test_acc:.4f}')")

img(doc,f'{IMG}/07_overfit_check.png',Inches(5.5),'Gambar 8.1 — Training vs Testing Accuracy')

doc.add_heading('8.1 Analisis per Model',level=2)
doc.add_paragraph(f'• Random Forest: Train Acc={R["Random Forest"]["train_acc"]}, Test Acc={R["Random Forest"]["test_acc"]}. Selisih cukup besar ({R["Random Forest"]["train_acc"]-R["Random Forest"]["test_acc"]:.4f}), mengindikasikan sedikit overfitting. Model menghafal data training terlalu detail.')
doc.add_paragraph(f'• Logistic Regression: Train Acc={R["Logistic Regression"]["train_acc"]}, Test Acc={R["Logistic Regression"]["test_acc"]}. Selisih kecil ({R["Logistic Regression"]["train_acc"]-R["Logistic Regression"]["test_acc"]:.4f}), menunjukkan generalisasi yang baik.')
doc.add_paragraph(f'• SVM: Train Acc={R["SVM"]["train_acc"]}, Test Acc={R["SVM"]["test_acc"]}. Selisih kecil ({R["SVM"]["train_acc"]-R["SVM"]["test_acc"]:.4f}), generalisasi cukup baik meskipun gagal mendeteksi kelas minoritas.')
doc.add_page_break()

# ===== 9. KESIMPULAN =====
doc.add_heading('9. Kesimpulan',level=1)
doc.add_paragraph('Berdasarkan analisis data modeling yang telah dilakukan, berikut kesimpulan yang diperoleh:')

doc.add_heading('9.1 Ringkasan Hasil',level=2)
doc.add_paragraph('1. Train/Test Split: Dataset dibagi 80:20 dengan stratified sampling. Terdapat 14 fitur dan 1 target variabel (depression_label). Dataset sangat imbalanced — kelas 0 (Tidak Depresi) mendominasi ~97.5% data.', style='List Number')
doc.add_paragraph('2. Random Forest: Akurasi 97.5% tetapi Precision/Recall/F1 = 0 untuk kelas positif. Model gagal mendeteksi satupun kasus depresi.', style='List Number')
doc.add_paragraph('3. Logistic Regression: Model terbaik dengan akurasi 98.75%, Precision 1.0, Recall 0.5, F1 0.6667. Berhasil mendeteksi 3 dari 6 kasus depresi tanpa false alarm.', style='List Number')
doc.add_paragraph('4. SVM: Akurasi 97.5% tetapi Precision/Recall/F1 = 0 — sama seperti Random Forest, gagal mendeteksi kasus depresi.', style='List Number')

doc.add_heading('9.2 Model Terbaik',level=2)
doc.add_paragraph('Logistic Regression dipilih sebagai model terbaik karena:\n• Akurasi tertinggi (98.75%)\n• Satu-satunya model yang berhasil mendeteksi kasus depresi (TP=3)\n• Tidak menghasilkan false positive (Precision=1.0)\n• Generalisasi yang baik (selisih train-test accuracy kecil)')

doc.add_heading('9.3 Feature Importance',level=2)
doc.add_paragraph('Dari analisis feature importance Random Forest, faktor yang paling berpengaruh terhadap depresi remaja adalah:\n• sleep_hours (20.99%) — jam tidur\n• daily_social_media_hours (19.19%) — penggunaan media sosial\n• stress_level (13.60%) — tingkat stres\n• anxiety_level (12.00%) — tingkat kecemasan')

doc.add_heading('9.4 Catatan Penting',level=2)
doc.add_paragraph('Dataset memiliki ketidakseimbangan kelas yang sangat ekstrem (97.5% kelas 0 vs 2.5% kelas 1). Hal ini menyebabkan sebagian besar model cenderung memprediksi semua data sebagai kelas mayoritas. Untuk pengembangan lebih lanjut, disarankan:\n• Menerapkan teknik resampling (SMOTE, oversampling, atau undersampling)\n• Menggunakan class_weight="balanced" pada parameter model\n• Menggunakan metrik evaluasi yang lebih sesuai untuk data imbalanced (F1, AUC-ROC)\n• Menerapkan K-Fold Cross Validation dan GridSearchCV untuk optimasi hyperparameter')

# SAVE
doc.save('laporan_data_modeling.docx')
print("laporan_data_modeling.docx berhasil dibuat!")
